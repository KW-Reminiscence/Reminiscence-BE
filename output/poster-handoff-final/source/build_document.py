"""Build the editable poster handoff DOCX from the final Markdown manuscript."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "Reminiscence_포스터_본문_최종본.md"
OUTPUT = ROOT / "Reminiscence_포스터_본문_최종본.docx"

FONT = "AppleGothic"
INK = RGBColor(31, 41, 51)
MUTED = RGBColor(102, 114, 126)
BLUE = RGBColor(46, 116, 181)
BLUE_DARK = RGBColor(31, 77, 120)
GOLD = RGBColor(176, 122, 27)
TABLE_FILL = "F4F6F9"
TABLE_BORDER = "B8C1CA"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 100, "bottom": 100, "start": 120, "end": 120}


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: RGBColor | None = None,
    weight: bool = False,
) -> None:
    run.font.name = FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:eastAsia"), FONT)
    r_fonts.set(qn("w:cs"), FONT)
    r_fonts.set(qn("w:hint"), "eastAsia")
    language = ensure_child(r_pr, "w:lang")
    language.set(qn("w:val"), "ko-KR")
    language.set(qn("w:eastAsia"), "ko-KR")
    run.bold = weight
    run.italic = False
    run.underline = False
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal_r_pr = normal._element.get_or_add_rPr()
    normal_r_fonts = normal_r_pr.get_or_add_rFonts()
    normal_r_fonts.set(qn("w:ascii"), FONT)
    normal_r_fonts.set(qn("w:hAnsi"), FONT)
    normal_r_fonts.set(qn("w:eastAsia"), FONT)
    normal_r_fonts.set(qn("w:cs"), FONT)
    normal_r_fonts.set(qn("w:hint"), "eastAsia")
    normal_language = ensure_child(normal_r_pr, "w:lang")
    normal_language.set(qn("w:val"), "ko-KR")
    normal_language.set(qn("w:eastAsia"), "ko-KR")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, BLUE_DARK, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        style.font.name = FONT
        style_r_pr = style._element.get_or_add_rPr()
        style_r_fonts = style_r_pr.get_or_add_rFonts()
        style_r_fonts.set(qn("w:ascii"), FONT)
        style_r_fonts.set(qn("w:hAnsi"), FONT)
        style_r_fonts.set(qn("w:eastAsia"), FONT)
        style_r_fonts.set(qn("w:cs"), FONT)
        style_r_fonts.set(qn("w:hint"), "eastAsia")
        style_language = ensure_child(style_r_pr, "w:lang")
        style_language.set(qn("w:val"), "ko-KR")
        style_language.set(qn("w:eastAsia"), "ko-KR")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = False
        style.font.italic = False
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.line_spacing = 1.15


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def field_run(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_header_footer(document: Document) -> None:
    for section in document.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("REMINISCENCE · POSTER HANDOFF")
        set_run_font(run, size=8.5, color=MUTED)

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        run = fp.add_run("KW-Reminiscence  |  ")
        set_run_font(run, size=8.5, color=MUTED)
        field_run(fp, "PAGE")


def add_cover(document: Document) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(96)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("POSTER HANDOFF · FINAL COPY")
    set_run_font(run, size=10.5, color=GOLD)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Reminiscence")
    set_run_font(run, size=30, color=BLUE_DARK)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(
        "가족사진 기반 회상 대화와 개인별 생활 패턴 변화 알림을 위한 스마트 케어 액자"
    )
    set_run_font(run, size=15, color=INK)

    english = document.add_paragraph()
    english.alignment = WD_ALIGN_PARAGRAPH.CENTER
    english.paragraph_format.space_after = Pt(56)
    run = english.add_run(
        "A Smart Care Frame for Photo-Elicited Conversation and "
        "Personalized Routine Change Monitoring"
    )
    set_run_font(run, size=11, color=MUTED)

    owner = document.add_paragraph()
    owner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    owner.paragraph_format.space_after = Pt(5)
    run = owner.add_run("KW-Reminiscence")
    set_run_font(run, size=11, color=BLUE_DARK)

    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run("2026-07-28")
    set_run_font(run, size=9.5, color=MUTED)

    page_break = document.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_width(parent, tag: str, width_dxa: int) -> None:
    width = ensure_child(parent, tag)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(width_dxa))


def apply_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("table widths must sum to 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    set_width(table_pr, "w:tblW", CONTENT_WIDTH_DXA)
    indent = ensure_child(table_pr, "w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    layout = ensure_child(table_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for column_index, width in enumerate(widths):
        table.columns[column_index].width = Twips(width)
    for row in table.rows:
        row.height = None
        for column_index, cell in enumerate(row.cells):
            width = widths[column_index]
            cell.width = Twips(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            set_width(tc_pr, "w:tcW", width)
            tc_mar = ensure_child(tc_pr, "w:tcMar")
            for side, value in CELL_MARGINS_DXA.items():
                margin = ensure_child(tc_mar, f"w:{side}")
                margin.set(qn("w:w"), str(value))
                margin.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = ensure_child(tc_pr, "w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = ensure_child(tc_pr, "w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = ensure_child(borders, f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), TABLE_BORDER)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def choose_table_widths(headers: list[str]) -> list[int]:
    if len(headers) == 2:
        return [2700, 6660]
    if headers and headers[0] == "데이터":
        return [2100, 3200, 4060]
    if headers and headers[0] == "검사":
        return [1650, 2800, 4910]
    if len(headers) == 3:
        return [2100, 3000, 4260]
    base = CONTENT_WIDTH_DXA // len(headers)
    widths = [base for _ in headers]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = choose_table_widths(rows[0])
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if row_index == 0 or column_index == 0
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = p.add_run(value)
            set_run_font(
                run,
                size=8.5 if row_index else 9,
                color=BLUE_DARK if row_index == 0 else INK,
            )
            set_cell_borders(cell)
            if row_index == 0:
                set_cell_shading(cell, TABLE_FILL)
    mark_header_row(table.rows[0])
    apply_table_geometry(table, widths)
    after = document.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)


def add_heading(document: Document, level: int, value: str) -> None:
    p = document.add_paragraph(style=f"Heading {min(level, 3)}")
    if level == 1 and value in {
        "A1 기술 및 시연 포스터 원고",
        "Figure 캡션 원고",
        "참고문헌",
    }:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt({1: 18, 2: 18, 3: 10}.get(level, 6))
    p.paragraph_format.space_after = Pt({1: 10, 2: 6, 3: 4}.get(level, 4))
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(value)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}.get(level, 11),
        color={1: BLUE, 2: BLUE, 3: BLUE_DARK}.get(level, INK),
    )


def add_body_paragraph(document: Document, value: str, *, size: float = 11) -> None:
    p = document.add_paragraph()
    p.paragraph_format.widow_control = True
    run = p.add_run(value.replace("`", ""))
    set_run_font(run, size=size, color=INK)


IMAGE_RE = re.compile(r"^!\[(?P<alt>.+?)\]\((?P<path>.+?)\)$")


def add_image(document: Document, relative_path: str, alt_text: str) -> None:
    image_path = ROOT / relative_path
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline_shape = run.add_picture(str(image_path), width=Inches(6.45))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def build_body(document: Document, source_text: str) -> None:
    lines = source_text.splitlines()
    start = lines.index("# A0 메인 포스터 원고")
    index = start
    paragraph_buffer: list[str] = []
    current_h1 = ""

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            size = 11
            if current_h1 == "Figure 캡션 원고":
                size = 10.5
            elif current_h1 == "참고문헌":
                size = 10
            add_body_paragraph(
                document,
                " ".join(part.strip() for part in paragraph_buffer),
                size=size,
            )
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            flush_paragraph()
            index += 1
            continue
        image_match = IMAGE_RE.match(line)
        if image_match:
            flush_paragraph()
            add_image(document, image_match.group("path"), image_match.group("alt"))
            index += 1
            continue
        if line.startswith("|"):
            flush_paragraph()
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, parse_table(table_lines))
            continue
        if line.startswith("#"):
            flush_paragraph()
            level = len(line) - len(line.lstrip("#"))
            heading = line[level:].strip()
            if level == 1:
                current_h1 = heading
            add_heading(document, level, heading)
            index += 1
            continue
        paragraph_buffer.append(line)
        index += 1
    flush_paragraph()


def main() -> None:
    document = Document()
    configure_section(document.sections[0])
    configure_styles(document)
    configure_header_footer(document)
    add_cover(document)
    build_body(document, SOURCE.read_text(encoding="utf-8"))
    document.core_properties.title = "Reminiscence 포스터 본문 최종본"
    document.core_properties.subject = "A0 메인 포스터와 A1 기술·시연 포스터 handoff"
    document.core_properties.author = "KW-Reminiscence"
    document.core_properties.keywords = "Reminiscence, poster, Raspberry Pi"
    document.core_properties.comments = (
        "narrative_proposal preset with Korean font and regular-weight overrides"
    )
    document.save(OUTPUT)


if __name__ == "__main__":
    main()
